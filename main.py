
import os
import PyPDF2
from docx import Document
import bijoy_converter  # Updated import

def extract_text_with_pypdf2(pdf_path):
    text_content = []
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text:
                unicode_text = bijoy_converter.convert_bijoy_to_unicode(text)  # Updated function call
                full_text = f"Page {page_num + 1}:\n" + unicode_text + "\n"
                text_content.append(full_text)
            else:
                text_content.append(f"Page {page_num + 1}: No text found.\n")
    return text_content

def pdf_to_word(pdf_path):
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    word_path = f"{base_name}.docx"

    doc = Document()
    

    text_content = extract_text_with_pypdf2(pdf_path)

    for page_num, page_text in enumerate(text_content):
        doc.add_heading(f"Page {page_num + 1}", level=2)
        doc.add_paragraph(page_text)
        doc.add_page_break()

    doc.save(word_path)
    print(f"✅ Conversion complete! The Word document is saved at: {word_path}")

# Example usage
pdf_to_word("bmcq.pdf")
